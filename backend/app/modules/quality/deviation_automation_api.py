"""偏差报告自动化 API"""

import logging
import uuid
import json
from pathlib import Path
from typing import Optional
from datetime import datetime, date

from fastapi import APIRouter, Depends, File, HTTPException, Query, UploadFile, Form, Body
from pydantic import BaseModel
from sqlalchemy import func, select, and_, or_
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.database import get_db

logger = logging.getLogger(__name__)


class ApiResponse(BaseModel):
    code: int = 200
    message: str = "Success"
    data: Optional[dict | list] = None


router = APIRouter(prefix="/deviation-automation", tags=["偏差报告自动化"])


# ============ SOP规则管理 API ============

class SOPRuleCreate(BaseModel):
    sop_code: str
    sop_full_name: str
    sop_version: str
    business_tag: Optional[str] = None
    standard_limit: Optional[str] = None
    standard_sentence: Optional[str] = None
    sop_file_path: Optional[str] = None


class SOPRuleUpdate(BaseModel):
    sop_code: Optional[str] = None
    sop_full_name: Optional[str] = None
    sop_version: Optional[str] = None
    business_tag: Optional[str] = None
    standard_limit: Optional[str] = None
    standard_sentence: Optional[str] = None
    status: Optional[int] = None
    sop_file_path: Optional[str] = None


@router.get("/sop-rules", summary="查询SOP规则列表")
async def list_sop_rules(
    keyword: Optional[str] = None,
    sop_code: Optional[str] = None,
    sop_full_name: Optional[str] = None,
    status: Optional[int] = None,
    page: int = 1,
    page_size: int = 20,
    db: AsyncSession = Depends(get_db),
):
    """查询SOP规则列表，支持模糊搜索和状态过滤"""
    conditions = []
    if keyword:
        conditions.append(
            or_(
                SOPRule.sop_code.ilike(f"%{keyword}%"),
                SOPRule.sop_full_name.ilike(f"%{keyword}%"),
                SOPRule.business_tag.ilike(f"%{keyword}%")
            )
        )
    if sop_code:
        conditions.append(SOPRule.sop_code.ilike(f"%{sop_code}%"))
    if sop_full_name:
        conditions.append(SOPRule.sop_full_name.ilike(f"%{sop_full_name}%"))
    if status is not None:
        conditions.append(SOPRule.status == status)

    query = select(SOPRule)
    if conditions:
        query = query.where(and_(*conditions))

    # 计数
    count_query = select(func.count()).select_from(query.subquery())
    count_result = await db.execute(count_query)
    total = count_result.scalar()

    # 分页
    query = query.order_by(SOPRule.create_time.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)

    result = await db.execute(query)
    items = result.scalars().all()

    rules = []
    for r in items:
        rules.append({
            "id": r.id,
            "sop_code": r.sop_code,
            "sop_full_name": r.sop_full_name,
            "sop_version": r.sop_version,
            "business_tag": r.business_tag,
            "standard_limit": r.standard_limit,
            "standard_sentence": r.standard_sentence,
            "sop_file_path": r.sop_file_path,
            "status": r.status,
            "create_time": r.create_time.isoformat() if r.create_time else None,
            "update_time": r.update_time.isoformat() if r.update_time else None,
        })

    return ApiResponse(data={"items": rules, "total": total})


@router.post("/sop-rules", summary="新增SOP规则")
async def create_sop_rule(
    rule: SOPRuleCreate,
    db: AsyncSession = Depends(get_db),
):
    """新增SOP规则"""
    # 检查编号唯一性
    result = await db.execute(
        select(SOPRule).where(SOPRule.sop_code == rule.sop_code)
    )
    existing = result.scalar_one_or_none()
    if existing:
        raise HTTPException(status_code=400, detail="该SOP编号已存在，请重新填写")

    new_rule = SOPRule(
        sop_code=rule.sop_code,
        sop_full_name=rule.sop_full_name,
        sop_version=rule.sop_version,
        business_tag=rule.business_tag,
        standard_limit=rule.standard_limit,
        standard_sentence=rule.standard_sentence,
        sop_file_path=rule.sop_file_path,
        status=1,
    )
    db.add(new_rule)
    await db.flush()
    await db.refresh(new_rule)

    return ApiResponse(message="新增成功", data={
        "id": new_rule.id,
        "sop_code": new_rule.sop_code,
    })


@router.post("/migrate-sop", summary="数据库迁移-添加SOP文件字段")
async def migrate_sop_file_path(
    db: AsyncSession = Depends(get_db),
):
    """添加 sop_file_path 字段到 sop_rule 表"""
    from sqlalchemy import text

    try:
        await db.execute(text("""
            ALTER TABLE quality.sop_rule
            ADD COLUMN IF NOT EXISTS sop_file_path VARCHAR(512)
        """))
        await db.commit()
        return ApiResponse(message="迁移成功，sop_file_path 字段已添加")
    except Exception as e:
        await db.rollback()
        raise HTTPException(status_code=500, detail=f"迁移失败: {str(e)}")


@router.put("/sop-rules/{rule_id}", summary="编辑SOP规则")
async def update_sop_rule(
    rule_id: int,
    rule: SOPRuleUpdate,
    db: AsyncSession = Depends(get_db),
):
    """编辑SOP规则"""
    result = await db.execute(select(SOPRule).where(SOPRule.id == rule_id))
    db_rule = result.scalar_one_or_none()
    if not db_rule:
        raise HTTPException(status_code=404, detail="规则不存在")

    # 检查编号唯一性（如果修改了编号）
    if rule.sop_code and rule.sop_code != db_rule.sop_code:
        result = await db.execute(
            select(SOPRule).where(SOPRule.sop_code == rule.sop_code)
        )
        existing = result.scalar_one_or_none()
        if existing:
            raise HTTPException(status_code=400, detail="该SOP编号已存在，请重新填写")

    # 更新字段
    update_data = rule.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        if value is not None:
            setattr(db_rule, key, value)

    db_rule.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="更新成功")


@router.put("/sop-rules/{rule_id}/status", summary="切换SOP规则状态")
async def toggle_sop_rule_status(
    rule_id: int,
    status: int = Body(...),
    db: AsyncSession = Depends(get_db),
):
    """切换SOP规则状态（生效/停用）"""
    result = await db.execute(select(SOPRule).where(SOPRule.id == rule_id))
    db_rule = result.scalar_one_or_none()
    if not db_rule:
        raise HTTPException(status_code=404, detail="规则不存在")

    db_rule.status = status
    db_rule.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="状态切换成功")


@router.delete("/sop-rules/{rule_id}", summary="删除SOP规则")
async def delete_sop_rule(
    rule_id: int,
    db: AsyncSession = Depends(get_db),
):
    """删除SOP规则"""
    result = await db.execute(select(SOPRule).where(SOPRule.id == rule_id))
    db_rule = result.scalar_one_or_none()
    if not db_rule:
        raise HTTPException(status_code=404, detail="规则不存在")

    await db.delete(db_rule)
    await db.flush()

    return ApiResponse(message="删除成功")


# ============ 辅助函数 ============

def _identify_sop_codes_in_text(text: str) -> list:
    """
    从文本中识别 SOP 编号
    匹配模式如: SOP-C003010-03, SOP-XXX-XXX, SOP-C003010 等
    """
    import re
    # 匹配 SOP 编号模式
    sop_pattern = r'SOP-[A-Z0-9]+(?:-[A-Z0-9]+)*'
    matches = re.findall(sop_pattern, text, re.IGNORECASE)
    # 去重并标准化
    unique_codes = list(set([code.upper() for code in matches]))
    return unique_codes


async def _load_sop_file_content(sop_code: str, db: AsyncSession, backend_dir: Path) -> str:
    """
    根据 SOP 编号读取对应的 SOP 文件内容
    支持带 SOP- 前缀和不带前缀的编号匹配
    """
    import mammoth

    # 标准化 SOP 编号（去掉 SOP- 前缀）
    normalized_code = sop_code.upper()
    if normalized_code.startswith('SOP-'):
        normalized_code = normalized_code[4:]  # 去掉 SOP- 前缀

    # 查询 SOP 规则 - 优先完整匹配，否则尝试去掉 SOP- 前缀
    result = await db.execute(
        select(SOPRule).where(
            (SOPRule.sop_code == normalized_code) |
            (SOPRule.sop_code == sop_code.upper()),
            SOPRule.status == 1
        )
    )
    sop_rule = result.scalar_one_or_none()

    # 如果没找到，尝试去掉 SOP- 前缀后再匹配
    if not sop_rule and sop_code.upper().startswith('SOP-'):
        code_without_prefix = sop_code[4:]
        result = await db.execute(
            select(SOPRule).where(
                SOPRule.sop_code == code_without_prefix,
                SOPRule.status == 1
            )
        )
        sop_rule = result.scalar_one_or_none()

    if not sop_rule or not sop_rule.sop_file_path:
        return None

    # 读取文件内容
    file_path = backend_dir / sop_rule.sop_file_path.lstrip("/")
    if not file_path.exists():
        return None

    try:
        with open(file_path, "rb") as f:
            text_result = mammoth.extract_raw_text(f)
            return text_result.value
    except Exception as e:
        logger.warning(f"读取SOP文件失败 {sop_code}: {e}")
        return None


# ============ SOP文件管理 API ============

@router.post("/sop-rules/{rule_id}/upload", summary="上传SOP文件")
async def upload_sop_file(
    rule_id: int,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """上传SOP文件"""
    import mammoth

    # 验证文件类型
    if not file.filename or not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="仅支持doc、docx格式文件")

    # 检查规则是否存在
    result = await db.execute(select(SOPRule).where(SOPRule.id == rule_id))
    db_rule = result.scalar_one_or_none()
    if not db_rule:
        raise HTTPException(status_code=404, detail="规则不存在")

    # 读取文件内容
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件超出大小限制，请压缩后重试")

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="文件为空或已损坏")

    # 保存文件
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "sop_files"
    uploads_dir.mkdir(parents=True, exist_ok=True)

    file_ext = Path(file.filename).suffix.lower()
    saved_filename = f"sop_{rule_id}_{uuid.uuid4().hex}{file_ext}"
    file_path = uploads_dir / saved_filename

    with open(file_path, "wb") as f:
        f.write(content)

    # 解析内容
    try:
        with open(file_path, "rb") as docx_file:
            text_result = mammoth.extract_raw_text(docx_file)
            plain_text = text_result.value
    except Exception:
        plain_text = ""

    # 更新规则记录
    relative_path = f"/uploads/sop_files/{saved_filename}"
    db_rule.sop_file_path = relative_path
    db_rule.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="上传成功", data={
        "rule_id": rule_id,
        "file_path": relative_path,
        "plain_text": plain_text,
    })


@router.post("/sop-rules/{rule_id}/ai-parse", summary="AI解析SOP文件")
async def ai_parse_sop_file(
    rule_id: int,
    db: AsyncSession = Depends(get_db),
):
    """AI解析SOP文件内容，提取关键字段"""
    import mammoth
    from app.platform.ai.minimax_util import get_vision_util

    # 获取规则
    result = await db.execute(select(SOPRule).where(SOPRule.id == rule_id))
    db_rule = result.scalar_one_or_none()
    if not db_rule:
        raise HTTPException(status_code=404, detail="规则不存在")

    if not db_rule.sop_file_path:
        raise HTTPException(status_code=400, detail="请先上传SOP文件")

    try:
        # 读取文件内容
        backend_dir = Path(__file__).resolve().parent.parent.parent.parent
        file_path = backend_dir / db_rule.sop_file_path.lstrip("/")

        if not file_path.exists():
            raise HTTPException(status_code=500, detail="SOP文件不存在")

        with open(file_path, "rb") as f:
            text_result = mammoth.extract_raw_text(f)
            sop_text = text_result.value

        # 调用AI解析
        prompt = f"""你是一个专业的SOP文档分析助手。请从以下SOP文档中提取关键信息：

SOP文档内容：
{sop_text}

请提取以下信息并以JSON格式输出：
{{
    "sop_code": "SOP编号（如已有则保持原样）",
    "sop_full_name": "SOP全称/标题",
    "sop_version": "版本号（如V1.0）",
    "business_tag": "业务标签（如偏差管理、质量控制等）",
    "standard_limit": "标准限度/关键控制点",
    "standard_sentence": "标准语句/关键流程描述"
}}

请只输出JSON格式，不要添加任何解释说明。
"""

        vision_util = get_vision_util()
        ai_result = await vision_util.recognize_text(
            prompt=prompt,
            max_tokens=4096,
        )

        # 尝试解析AI返回的JSON
        parsed_data = None
        try:
            # 提取JSON部分
            import re
            json_match = re.search(r'\{[^}]+\}', ai_result, re.DOTALL)
            if json_match:
                parsed_data = json.loads(json_match.group())
        except Exception:
            pass

        # 如果解析成功，更新规则
        if parsed_data:
            if parsed_data.get("sop_code"):
                db_rule.sop_code = parsed_data["sop_code"]
            if parsed_data.get("sop_full_name"):
                db_rule.sop_full_name = parsed_data["sop_full_name"]
            if parsed_data.get("sop_version"):
                db_rule.sop_version = parsed_data["sop_version"]
            if parsed_data.get("business_tag"):
                db_rule.business_tag = parsed_data["business_tag"]
            if parsed_data.get("standard_limit"):
                db_rule.standard_limit = parsed_data["standard_limit"]
            if parsed_data.get("standard_sentence"):
                db_rule.standard_sentence = parsed_data["standard_sentence"]

        db_rule.update_time = datetime.now()
        await db.flush()

        return ApiResponse(message="AI解析完成", data={
            "rule_id": rule_id,
            "ai_result": ai_result,
            "parsed_data": parsed_data,
        })

    except Exception as e:
        # 检查是否是唯一约束冲突
        error_str = str(e)
        if "UniqueViolationError" in error_str or "sop_rule_sop_code_key" in error_str:
            # sop_code重复，跳过code更新，只更新其他字段
            if parsed_data and "sop_code" in parsed_data:
                logger.warning(f"sop_code {parsed_data['sop_code']} 已被其他记录使用，跳过code更新")
                # 需要先回滚当前事务，然后重新开始
                await db.rollback()
                # 重新获取规则
                result = await db.execute(select(SOPRule).where(SOPRule.id == rule_id))
                db_rule = result.scalar_one()
                # 更新其他字段
                if parsed_data.get("sop_full_name"):
                    db_rule.sop_full_name = parsed_data["sop_full_name"]
                if parsed_data.get("sop_version"):
                    db_rule.sop_version = parsed_data["sop_version"]
                if parsed_data.get("business_tag"):
                    db_rule.business_tag = parsed_data["business_tag"]
                if parsed_data.get("standard_limit"):
                    db_rule.standard_limit = parsed_data["standard_limit"]
                if parsed_data.get("standard_sentence"):
                    db_rule.standard_sentence = parsed_data["standard_sentence"]
                db_rule.update_time = datetime.now()
                await db.commit()
                return ApiResponse(message="AI解析完成(sop_code已存在，跳过更新)", data={
                    "rule_id": rule_id,
                    "ai_result": ai_result,
                    "parsed_data": parsed_data,
                    "warning": f"sop_code {parsed_data['sop_code']} 已被其他记录使用",
                })
        import traceback
        logger.error(f"AI解析SOP文件失败: {e}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"AI解析失败: {str(e)}")


@router.get("/sop-rules/{rule_id}/download", summary="下载SOP原文件")
async def download_sop_file(
    rule_id: int,
    db: AsyncSession = Depends(get_db),
):
    """下载SOP原文件"""
    from fastapi.responses import FileResponse

    result = await db.execute(select(SOPRule).where(SOPRule.id == rule_id))
    db_rule = result.scalar_one_or_none()
    if not db_rule:
        raise HTTPException(status_code=404, detail="规则不存在")

    if not db_rule.sop_file_path:
        raise HTTPException(status_code=404, detail="SOP文件不存在")

    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    file_path = backend_dir / db_rule.sop_file_path.lstrip("/")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=str(file_path),
        filename=f"{db_rule.sop_code}_{db_rule.sop_full_name[:20] if db_rule.sop_full_name else 'SOP'}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


# ============ 偏差任务管理 API ============

class DevTaskCreate(BaseModel):
    deviation_no: str
    creator: str
    auditor: Optional[str] = None
    report_date: date


@router.post("/tasks", summary="新建偏差任务")
async def create_dev_task(
    task: DevTaskCreate,
    db: AsyncSession = Depends(get_db),
):
    """新建偏差报告任务"""
    # 检查编号唯一性
    result = await db.execute(
        select(DevTask).where(DevTask.deviation_no == task.deviation_no)
    )
    existing = result.scalar_one_or_none()
    if existing:
        raise HTTPException(status_code=400, detail="该偏差编号已存在，请重新填写")

    new_task = DevTask(
        deviation_no=task.deviation_no,
        creator=task.creator,
        auditor=task.auditor,
        report_date=task.report_date,
        task_status=1,  # 待处理
    )
    db.add(new_task)
    await db.flush()
    await db.refresh(new_task)

    return ApiResponse(message="创建成功", data={
        "task_id": new_task.task_id,
        "deviation_no": new_task.deviation_no,
    })


@router.get("/tasks", summary="查询偏差任务列表")
async def list_dev_tasks(
    deviation_no: Optional[str] = None,
    creator: Optional[str] = None,
    task_status: Optional[int] = None,
    start_date: Optional[date] = None,
    end_date: Optional[date] = None,
    page: int = 1,
    page_size: int = 20,
    db: AsyncSession = Depends(get_db),
):
    """查询偏差任务列表"""
    conditions = []
    if deviation_no:
        conditions.append(DevTask.deviation_no.ilike(f"%{deviation_no}%"))
    if creator:
        conditions.append(DevTask.creator.ilike(f"%{creator}%"))
    if task_status is not None:
        conditions.append(DevTask.task_status == task_status)
    if start_date:
        conditions.append(DevTask.report_date >= start_date)
    if end_date:
        conditions.append(DevTask.report_date <= end_date)

    query = select(DevTask)
    if conditions:
        query = query.where(and_(*conditions))

    # 计数
    count_query = select(func.count()).select_from(query.subquery())
    count_result = await db.execute(count_query)
    total = count_result.scalar()

    # 分页
    query = query.order_by(DevTask.create_time.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)

    result = await db.execute(query)
    items = result.scalars().all()

    tasks = []
    for t in items:
        tasks.append({
            "task_id": t.task_id,
            "deviation_no": t.deviation_no,
            "creator": t.creator,
            "auditor": t.auditor,
            "report_date": t.report_date.isoformat() if t.report_date else None,
            "original_file_path": t.original_file_path,
            "standard_file_path": t.standard_file_path,
            "task_status": t.task_status,
            "create_time": t.create_time.isoformat() if t.create_time else None,
            "update_time": t.update_time.isoformat() if t.update_time else None,
        })

    return ApiResponse(data={"items": tasks, "total": total})


@router.get("/tasks/{task_id}", summary="获取任务详情")
async def get_dev_task(
    task_id: int,
    db: AsyncSession = Depends(get_db),
):
    """获取任务详情"""
    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    return ApiResponse(data={
        "task_id": task.task_id,
        "deviation_no": task.deviation_no,
        "creator": task.creator,
        "auditor": task.auditor,
        "report_date": task.report_date.isoformat() if task.report_date else None,
        "original_file_path": task.original_file_path,
        "standard_file_path": task.standard_file_path,
        "task_status": task.task_status,
        "ai_result": task.ai_result,
        "create_time": task.create_time.isoformat() if task.create_time else None,
        "update_time": task.update_time.isoformat() if task.update_time else None,
    })


@router.put("/tasks/{task_id}/status", summary="更新任务状态")
async def update_task_status(
    task_id: int,
    status: int,
    db: AsyncSession = Depends(get_db),
):
    """更新任务状态"""
    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    task.task_status = status
    task.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="状态更新成功")


class AIResultUpdate(BaseModel):
    ai_result: str


@router.put("/tasks/{task_id}/update-ai-result", summary="更新AI处理结果")
async def update_ai_result(
    task_id: int,
    data: AIResultUpdate,
    db: AsyncSession = Depends(get_db),
):
    """更新任务的AI处理结果"""
    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    # 检查内容格式
    content = data.ai_result.strip()

    # 如果是 HTML 格式，自动转换为 JSON
    if (content.startswith('<!DOCTYPE') or
        content.startswith('<html') or
        content.startswith('<meta') or
        content.startswith('<head') or
        content.startswith('<body')):

        # 从 HTML 中提取纯文本并调用 AI 重新处理
        try:
            # 简单提取 HTML 中的文本内容
            import re
            text_content = re.sub(r'<[^>]+>', ' ', content)
            text_content = re.sub(r'\s+', ' ', text_content).strip()

            # 获取 SOP 规则
            sop_result = await db.execute(select(SOPRule).where(SOPRule.status == 1))
            sop_rules = sop_result.scalars().all()
            sop_context = "\n\n".join([
                f"[{r.sop_code}] {r.sop_full_name}\n版本: {r.sop_version}\n业务标签: {r.business_tag or '无'}\n标准限值: {r.standard_limit or '无'}\n标准语句: {r.standard_sentence or '无'}\n"
                for r in sop_rules
            ])

            prompt = f"""<|im_end|>
<|im_start|>user
你是一个专业的偏差报告处理助手。请将以下已编辑的报告内容重新整理为标准JSON格式。

【关键要求】
- 不要输出任何思考过程、推理步骤或解释说明
- 不要使用<think>标签
- 不要使用省略号（...）或任何截断符号
- 所有文本内容必须完整输出，不要省略
- 直接输出JSON格式结果

SOP规则库：
{sop_context}

已编辑的报告内容：
{text_content[:8000]}

请严格按照以下JSON格式输出，直接输出JSON：
{{
    "偏差描述": "提取的偏差描述内容",
    "根本原因": "提取的根本原因",
    "纠正预防措施": ["措施1", "措施2", "措施3", "措施4"],
    "检测方法": "提取的检测方法",
    "溶液制备过程": "提取的溶液制备过程",
    "检测结果": "提取的检测结果",
    "初步原因分析": "提取的初步原因分析",
    "深入调查": "提取的深入调查内容",
    "对目前批次产品质量的影响": "提取的内容",
    "对相邻批次质量的影响": "提取的内容",
    "对其他产品质量或系统的影响": "提取的内容",
    "历史偏差": "回顾过去6个月内的重复性偏差情况",
    "其他影响": "其他方面影响",
    "调查结论": "提取的调查结论",
    "根本原因类别": ["选中的原因类别"],
    "CA_PA措施": [
        {{"措施": "措施1内容", "说明": "说明1"}},
        {{"措施": "措施2内容", "说明": "说明2"}},
        {{"措施": "措施3内容", "说明": "说明3"}},
        {{"措施": "措施4内容", "说明": "说明4"}}
    ]
}}
直接输出JSON：<|im_end|>
<|im_start|>assistant
"""

            from app.platform.ai.minimax_util import get_vision_util
            vision_util = get_vision_util()
            ai_result = await vision_util.recognize_text(prompt=prompt, max_tokens=8192)

            # 解析并验证
            ai_data = _parse_ai_result(ai_result)
            logger.info(f"HTML内容已自动转换为JSON，包含 {len(ai_data)} 个字段")

            task.ai_result = ai_result
            task.update_time = datetime.now()
            await db.flush()
            return ApiResponse(message="保存成功（已自动转换为标准格式）")

        except Exception as e:
            logger.error(f"HTML转JSON失败: {e}")
            # 转换失败时保留原始内容
            task.ai_result = data.ai_result
            task.update_time = datetime.now()
            await db.flush()
            return ApiResponse(message="保存成功（格式未转换）")

    # 如果是 JSON 格式，直接保存
    task.ai_result = data.ai_result
    task.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="保存成功")


# ============ Word文件上传与解析 ============

@router.post("/upload", summary="上传并解析Word文件")
async def upload_and_parse_word(
    task_id: int = Form(...),
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """上传Word文件并解析内容"""
    import mammoth
    import shutil

    # 验证文件类型
    if not file.filename or not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="仅支持doc、docx格式文件")

    # 验证文件大小 (20MB)
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="文件超出大小限制，请压缩后重试")

    if len(content) == 0:
        raise HTTPException(status_code=400, detail="文件为空或已损坏，请检查文件后重新上传")

    # 检查任务是否存在
    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    # 保存文件
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "deviation_automation"
    uploads_dir.mkdir(parents=True, exist_ok=True)

    file_ext = Path(file.filename).suffix.lower()
    saved_filename = f"dev_{task_id}_{uuid.uuid4().hex}{file_ext}"
    file_path = uploads_dir / saved_filename

    with open(file_path, "wb") as f:
        f.write(content)

    # 解析内容
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html_content = result.value
            messages = [msg.message for msg in result.messages]
    except Exception as e:
        raise HTTPException(status_code=400, detail="文件无法解析，请检查文件后重新上传")

    # 提取纯文本
    try:
        with open(file_path, "rb") as docx_file:
            text_result = mammoth.extract_raw_text(docx_file)
            plain_text = text_result.value
    except Exception:
        plain_text = ""

    # 更新任务记录
    relative_path = f"/uploads/deviation_automation/{saved_filename}"
    task.original_file_path = relative_path
    task.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="上传成功", data={
        "task_id": task_id,
        "file_path": relative_path,
        "html_content": html_content,
        "plain_text": plain_text,
        "warnings": messages,
    })


# ============ AI处理 ============

# 标准报告结构模板 - 用于指导AI提取和结构化输出
STANDARD_REPORT_TEMPLATE = """标准偏差报告包含以下7个固定部分：
1. 总结概述 (Executive Summary)
   - 1.1 偏差描述
   - 1.2 根本原因
   - 1.3 纠正预防措施（最多4项）
2. 根本原因调查 (Root Cause Investigation)
   - 2.1 背景介绍
     - 检测方法
     - 溶液制备过程
     - 检测结果
   - 2.2 根本原因分析及调查
   - 2.3 深入调查
3. 影响评估 (Impact Assessment)
   - 对目前批次产品质量的影响
   - 对偏差产品相邻批次质量的影响
   - 对其他产品质量或系统的影响
   - 回顾过去6个月内的重复性偏差发生情况
   - 其他方面影响
4. 调查结论 (Investigation Conclusion)
   - 调查结论内容
   - 根本原因类别（多选）
5. 纠正预防措施 (CA PA)
   - 最多4项措施，每项包含措施内容和说明
6. 签名区
   - 调查人/日期、QA调查人/日期"""


@router.post("/tasks/{task_id}/ai-process", summary="触发AI处理")
async def trigger_ai_process(
    task_id: int,
    db: AsyncSession = Depends(get_db),
):
    """触发AI处理原始Word内容 - 提取结构化数据"""
    import mammoth
    from app.platform.ai.minimax_util import get_vision_util

    # 获取任务
    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    if not task.original_file_path:
        raise HTTPException(status_code=400, detail="请先上传原始文件")

    # 更新状态为AI处理中
    task.task_status = 2  # AI处理中
    task.update_time = datetime.now()
    await db.flush()

    try:
        # 读取原始文件内容
        backend_dir = Path(__file__).resolve().parent.parent.parent.parent
        original_path = backend_dir / task.original_file_path.lstrip("/")

        if not original_path.exists():
            raise HTTPException(status_code=500, detail="原始文件不存在")

        with open(original_path, "rb") as f:
            text_result = mammoth.extract_raw_text(f)
            original_text = text_result.value

        # 获取生效的SOP规则
        sop_result = await db.execute(
            select(SOPRule).where(SOPRule.status == 1)
        )
        sop_rules = sop_result.scalars().all()

        # 构建 SOP 规则元数据上下文
        sop_context = "\n\n".join([
            f"[{r.sop_code}] {r.sop_full_name}\n"
            f"版本: {r.sop_version}\n"
            f"业务标签: {r.business_tag or '无'}\n"
            f"标准限值: {r.standard_limit or '无'}\n"
            f"标准语句: {r.standard_sentence or '无'}\n"
            for r in sop_rules
        ])

        # ========== 增强：自动识别并读取相关 SOP 文件 ==========
        # 从原始报告中识别涉及的 SOP 编号
        identified_sop_codes = _identify_sop_codes_in_text(original_text)
        logger.info(f"报告中识别的SOP编号: {identified_sop_codes}")

        # 读取识别到的 SOP 文件内容
        sop_files_context = ""
        if identified_sop_codes:
            sop_files_context = "\n\n" + "="*60 + "\n"
            sop_files_context += "相关 SOP 文件内容（用于参考标准操作规程）\n"
            sop_files_context += "="*60 + "\n\n"

            for sop_code in identified_sop_codes:
                sop_content = await _load_sop_file_content(sop_code, db, backend_dir)
                if sop_content:
                    sop_files_context += f"【{sop_code}】\n{sop_content[:5000]}\n\n"  # 限制每个 SOP 内容长度
                    logger.info(f"成功加载 SOP 文件: {sop_code}")
                else:
                    logger.warning(f"未找到匹配的 SOP 文件: {sop_code}")
                    sop_files_context += f"【{sop_code}】\n[文件未找到]\n\n"

        # 调用AI - 按照标准模板结构提取数据
        # 使用特殊标记分隔用户输入和AI输出区域
        prompt = f"""<|im_end|>
<|im_start|>user
你是一个专业的偏差报告处理助手。任务是：
1. 从用户上传的原始偏差报告中提取关键信息
2. 根据SOP规则库和相关SOP文件内容对内容进行标准化
3. 直接输出JSON格式的结构化数据

【关键要求】
- 不要输出任何思考过程、推理步骤或解释说明
- 不要使用<think>标签
- 不要使用省略号（...）或任何截断符号
- 所有文本内容必须完整输出，不要省略
- 直接输出JSON格式结果
- 如果报告中引用了相关SOP，应结合SOP文件中的标准要求进行内容验证和优化

SOP规则库：
{sop_context}
{sop_files_context}
标准报告结构模板：
{STANDARD_REPORT_TEMPLATE}

原始偏差报告内容：
{original_text}

请严格按照以下JSON格式输出，直接输出JSON：
{{
    "偏差描述": "提取的偏差描述内容",
    "根本原因": "提取的根本原因",
    "纠正预防措施": ["措施1", "措施2", "措施3", "措施4"],
    "检测方法": "提取的检测方法",
    "溶液制备过程": "提取的溶液制备过程",
    "检测结果": "提取的检测结果",
    "初步原因分析": "提取的初步原因分析",
    "深入调查": "提取的深入调查内容",
    "对目前批次产品质量的影响": "提取的内容",
    "对相邻批次质量的影响": "提取的内容",
    "对其他产品质量或系统的影响": "提取的内容",
    "历史偏差": "回顾过去6个月内的重复性偏差情况",
    "其他影响": "其他方面影响",
    "调查结论": "提取的调查结论",
    "根本原因类别": ["选中的原因类别"],
    "CA_PA措施": [
        {{"措施": "措施1内容", "说明": "说明1"}},
        {{"措施": "措施2内容", "说明": "说明2"}},
        {{"措施": "措施3内容", "说明": "说明3"}},
        {{"措施": "措施4内容", "说明": "说明4"}}
    ]
}}

如果原始报告中某项内容缺失，字段值留空字符串或空数组。
直接输出JSON格式的结果，不要有其他内容：<|im_end|>
<|im_start|>assistant
"""

        vision_util = get_vision_util()
        ai_result = await vision_util.recognize_text(
            prompt=prompt,
            max_tokens=8192,
        )

        # 保存AI结果（结构化JSON数据）
        task.ai_result = ai_result
        task.task_status = 3  # 已生成
        task.update_time = datetime.now()
        await db.flush()

        return ApiResponse(message="AI处理完成", data={
            "task_id": task_id,
            "ai_result": ai_result,
            "task_status": task.task_status,
        })

    except Exception as e:
        logger.error(f"AI处理失败: {e}")
        task.task_status = 1  # 恢复待处理状态
        task.update_time = datetime.now()
        await db.flush()
        raise HTTPException(status_code=500, detail=f"AI解析失败，请稍后重试")


def _parse_ai_result(ai_result: str) -> dict:
    """解析AI返回的JSON数据"""
    import re
    import json as json_lib
    
    # 1. 移除AI思考过程标签
    raw_text = re.sub(r'<think>[\s\S]*?</think>', '', ai_result)

    # 2. 移除markdown代码块标记
    raw_text = re.sub(r'^```json\s*', '', raw_text, flags=re.MULTILINE)
    raw_text = re.sub(r'^```\s*$', '', raw_text, flags=re.MULTILINE)

    # 3. 查找AI响应的开始位置（<|im_start|>assistant 之后）
    # 这处理了使用特殊标记分隔的对话格式
    assistant_pos = raw_text.find('<|im_start|>assistant')
    if assistant_pos != -1:
        raw_text = raw_text[assistant_pos + len('<|im_start|>assistant'):]
    else:
        # 如果没有 assistant 标记，移除 <|im_end|> 之后的文本
        im_end_pos = raw_text.find('<|im_end|>')
        if im_end_pos != -1:
            raw_text = raw_text[im_end_pos + len('<|im_end|>'):]

    # 4. 查找JSON开始位置（第一个{或[）
    json_start = -1
    for i, c in enumerate(raw_text):
        if c == '{' or c == '[':
            json_start = i
            break
    
    if json_start == -1:
        raise ValueError("未找到JSON开始标记")
    
    json_text = raw_text[json_start:].strip()

    # 5. 尝试找到正确的JSON结束位置
    brace_count = 0
    bracket_count = 0
    in_string = False
    escape_next = False
    json_end = 0
    
    for i, c in enumerate(json_text):
        if escape_next:
            escape_next = False
            continue
        if c == '\\':
            escape_next = True
            continue
        if c == '"' and not escape_next:
            in_string = not in_string
            continue
        if in_string:
            continue
            
        if c == '{':
            brace_count += 1
        elif c == '}':
            brace_count -= 1
            if brace_count == 0 and bracket_count == 0:
                json_end = i + 1
                break
        elif c == '[':
            bracket_count += 1
        elif c == ']':
            bracket_count -= 1
            if brace_count == 0 and bracket_count == 0:
                json_end = i + 1
                break
    
    if json_end > 0:
        json_text = json_text[:json_end]

    # 6. 清理可能的尾部干扰字符
    json_text = json_text.strip()
    if not json_text.endswith(('}', ']')):
        last_brace = json_text.rfind('}')
        last_bracket = json_text.rfind(']')
        end_pos = max(last_brace, last_bracket)
        if end_pos > 0:
            json_text = json_text[:end_pos + 1]
    
    return json_lib.loads(json_text)


def _generate_preview_html(ai_data: dict, task_info: dict = None) -> str:
    """根据AI数据生成HTML预览内容"""
    
    def safe_get(key: str, default: str = ''):
        value = ai_data.get(key, default)
        if isinstance(value, list):
            return value
        return str(value) if value else default
    
    # 获取数据
    deviation_desc = safe_get('偏差描述')
    root_cause = safe_get('根本原因')
    measures = safe_get('纠正预防措施', [])
    detection_method = safe_get('检测方法')
    solution_prep = safe_get('溶液制备过程')
    detection_result = safe_get('检测结果')
    initial_analysis = safe_get('初步原因分析')
    deep_investigation = safe_get('深入调查')
    impact_current = safe_get('对目前批次产品质量的影响')
    impact_adjacent = safe_get('对相邻批次质量的影响')
    impact_other = safe_get('对其他产品质量或系统的影响')
    historical_dev = safe_get('历史偏差')
    other_impact = safe_get('其他影响')
    conclusion = safe_get('调查结论')
    root_cause_category = safe_get('根本原因类别', [])
    capa_measures = safe_get('CA_PA措施', [])
    
    # 格式化措施列表
    measures_html = ''
    if isinstance(measures, list) and len(measures) > 0:
        for i, m in enumerate(measures, 1):
            if m:
                measures_html += f'<p style="margin: 4px 0; text-indent: 2em;">{i}. {m}</p>'
    else:
        measures_html = '<p style="color: #999;">无</p>'
    
    # 格式化CA/PA措施表格
    capa_html = ''
    if isinstance(capa_measures, list) and len(capa_measures) > 0:
        capa_html = '''
        <table style="width: 100%; border-collapse: collapse; margin-top: 8px;">
            <thead>
                <tr style="background: #f0f0f0;">
                    <th style="border: 1px solid #ddd; padding: 8px; width: 50px;">序号</th>
                    <th style="border: 1px solid #ddd; padding: 8px;">纠正预防措施</th>
                    <th style="border: 1px solid #ddd; padding: 8px;">说明</th>
                </tr>
            </thead>
            <tbody>
        '''
        for i, item in enumerate(capa_measures, 1):
            if isinstance(item, dict):
                measure_text = item.get('措施', '')
                desc_text = item.get('说明', '')
            else:
                measure_text = str(item)
                desc_text = ''
            capa_html += f'''
                <tr>
                    <td style="border: 1px solid #ddd; padding: 8px; text-align: center;">{i}</td>
                    <td style="border: 1px solid #ddd; padding: 8px;">{measure_text or '无'}</td>
                    <td style="border: 1px solid #ddd; padding: 8px;">{desc_text or '无'}</td>
                </tr>
            '''
        capa_html += '</tbody></table>'
    else:
        capa_html = '<p style="color: #999;">暂无纠正预防措施</p>'
    
    # 根本原因类别
    category_html = ''
    if isinstance(root_cause_category, list) and len(root_cause_category) > 0:
        category_html = '、'.join(root_cause_category)
    else:
        category_html = '无'
    
    # 基础信息
    deviation_no = task_info.get('deviation_no', '') if task_info else ''
    creator = task_info.get('creator', '') if task_info else ''
    auditor = task_info.get('auditor', '') or '无' if task_info else '无'
    report_date = task_info.get('report_date', '') if task_info else ''
    
    html = f'''
<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
    body {{
        font-family: 'Times New Roman', 'SimSun', serif;
        font-size: 12pt;
        line-height: 1.5;
        margin: 0;
        padding: 20px;
        color: #333;
    }}
    .report-title {{
        text-align: center;
        font-size: 20pt;
        font-weight: bold;
        margin-bottom: 24px;
    }}
    .info-table {{
        width: 100%;
        border-collapse: collapse;
        margin-bottom: 24px;
    }}
    .info-table td {{
        border: 1px solid #ddd;
        padding: 8px;
    }}
    .info-table .label {{
        background: #f5f5f5;
        font-weight: bold;
        width: 25%;
    }}
    h2 {{
        font-size: 14pt;
        font-weight: bold;
        margin-top: 20px;
        margin-bottom: 10px;
        border-bottom: 1px solid #333;
        padding-bottom: 4px;
    }}
    h3 {{
        font-size: 12pt;
        font-weight: bold;
        margin-top: 16px;
        margin-bottom: 8px;
    }}
    p {{
        margin: 8px 0;
        text-align: justify;
    }}
    .section {{
        margin-bottom: 16px;
    }}
    .signature-area {{
        margin-top: 40px;
    }}
    .signature-table {{
        width: 100%;
    }}
    .signature-table td {{
        padding: 12px 20px;
    }}
    @media print {{
        body {{ padding: 0; }}
    }}
</style>
</head>
<body>
    <h1 class="report-title">偏差调查报告</h1>
    
    <table class="info-table">
        <tr>
            <td class="label">偏差编号</td>
            <td>{deviation_no}</td>
            <td class="label">报告日期</td>
            <td>{report_date}</td>
        </tr>
        <tr>
            <td class="label">调查人</td>
            <td>{creator}</td>
            <td class="label">QA审核人</td>
            <td>{auditor}</td>
        </tr>
    </table>

    <h2>一、总结概述</h2>
    
    <div class="section">
        <h3>1.1 偏差描述</h3>
        <p style="text-indent: 2em;">{deviation_desc or '无'}</p>
    </div>
    
    <div class="section">
        <h3>1.2 根本原因</h3>
        <p style="text-indent: 2em;">{root_cause or '无'}</p>
    </div>
    
    <div class="section">
        <h3>1.3 纠正预防措施</h3>
        {measures_html}
    </div>

    <h2>二、根本原因调查</h2>
    
    <div class="section">
        <h3>2.1 背景介绍</h3>
        <p><strong>检测方法：</strong>{detection_method or '无'}</p>
        <p><strong>溶液制备过程：</strong>{solution_prep or '无'}</p>
        <p><strong>检测结果：</strong>{detection_result or '无'}</p>
    </div>
    
    <div class="section">
        <h3>2.2 根本原因分析及调查</h3>
        <p style="text-indent: 2em;">{initial_analysis or '无'}</p>
    </div>
    
    <div class="section">
        <h3>2.3 深入调查</h3>
        <p style="text-indent: 2em;">{deep_investigation or '无'}</p>
    </div>

    <h2>三、影响评估</h2>
    
    <div class="section">
        <p><strong>对目前批次产品质量的影响：</strong>{impact_current or '无'}</p>
        <p><strong>对偏差产品相邻批次质量的影响：</strong>{impact_adjacent or '无'}</p>
        <p><strong>对其他产品质量或系统的影响：</strong>{impact_other or '无'}</p>
        <p><strong>回顾过去6个月内的重复性偏差发生情况：</strong>{historical_dev or '无'}</p>
        <p><strong>其他方面影响：</strong>{other_impact or '无'}</p>
    </div>

    <h2>四、调查结论</h2>
    
    <div class="section">
        <p><strong>调查结论：</strong>{conclusion or '无'}</p>
        <p><strong>根本原因类别：</strong>{category_html}</p>
    </div>

    <h2>五、纠正预防措施</h2>
    
    <div class="section">
        {capa_html}
    </div>

    <h2>六、签名区</h2>
    
    <div class="signature-area">
        <table class="signature-table">
            <tr>
                <td>调查人签名：__________</td>
                <td>日期：__________</td>
            </tr>
            <tr>
                <td>QA调查人签名：__________</td>
                <td>日期：__________</td>
            </tr>
        </table>
    </div>
</body>
</html>
    '''
    return html


@router.get("/tasks/{task_id}/preview", summary="获取报告预览内容")
async def get_task_preview(
    task_id: int,
    db: AsyncSession = Depends(get_db),
):
    """获取填充后的报告预览内容（HTML格式）"""
    import html as html_lib

    # 获取任务
    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    if not task.ai_result:
        raise HTTPException(status_code=400, detail="请先进行AI处理")

    try:
        # 检查是否已经是HTML格式（contenteditable编辑后可能以<meta>开头）
        ai_result_stripped = task.ai_result.strip()
        if (ai_result_stripped.startswith('<!DOCTYPE') or
            ai_result_stripped.startswith('<html') or
            ai_result_stripped.startswith('<meta') or
            ai_result_stripped.startswith('<head') or
            ai_result_stripped.startswith('<body')):
            # 已经是HTML格式，直接返回
            return ApiResponse(data={
                'task_id': task_id,
                'html_content': html_lib.escape(task.ai_result),
                'plain_content': task.ai_result,
            })

        # 解析AI数据（JSON格式）
        try:
            ai_data = _parse_ai_result(task.ai_result)
        except Exception as parse_error:
            # JSON解析失败，尝试自动恢复
            logger.warning(f"JSON解析失败，尝试自动恢复: {parse_error}")
            try:
                # 重新读取原始文件并调用AI处理
                backend_dir = Path(__file__).resolve().parent.parent.parent.parent

                if not task.original_file_path:
                    raise HTTPException(status_code=400, detail="原始文件不存在，无法自动恢复")

                original_path = backend_dir / task.original_file_path.lstrip("/")
                if not original_path.exists():
                    raise HTTPException(status_code=500, detail="原始文件不存在")

                with open(original_path, "rb") as f:
                    text_result = mammoth.extract_raw_text(f)
                    original_text = text_result.value

                # 获取SOP规则
                sop_result = await db.execute(select(SOPRule).where(SOPRule.status == 1))
                sop_rules = sop_result.scalars().all()
                sop_context = "\n\n".join([
                    f"[{r.sop_code}] {r.sop_full_name}\n版本: {r.sop_version}\n业务标签: {r.business_tag or '无'}\n标准限值: {r.standard_limit or '无'}\n标准语句: {r.standard_sentence or '无'}\n"
                    for r in sop_rules
                ])

                # 识别并加载SOP文件
                identified_sop_codes = _identify_sop_codes_in_text(original_text)
                sop_files_context = ""
                if identified_sop_codes:
                    sop_files_context = "\n\n" + "="*60 + "\n相关 SOP 文件内容：\n" + "="*60 + "\n\n"
                    for sop_code in identified_sop_codes:
                        sop_content = await _load_sop_file_content(sop_code, db, backend_dir)
                        if sop_content:
                            sop_files_context += f"【{sop_code}】\n{sop_content[:5000]}\n\n"

                prompt = f"""<|im_end|>
<|im_start|>user
你是一个专业的偏差报告处理助手。任务是：
1. 从用户上传的原始偏差报告中提取关键信息
2. 根据SOP规则库和相关SOP文件内容对内容进行标准化
3. 直接输出JSON格式的结构化数据

【关键要求】
- 不要输出任何思考过程、推理步骤或解释说明
- 不要使用<think>标签
- 不要使用省略号（...）或任何截断符号
- 所有文本内容必须完整输出，不要省略
- 直接输出JSON格式结果

SOP规则库：
{sop_context}
{sop_files_context}
标准报告结构模板：
{STANDARD_REPORT_TEMPLATE}

原始偏差报告内容：
{original_text}

请严格按照以下JSON格式输出，直接输出JSON：
{{
    "偏差描述": "提取的偏差描述内容",
    "根本原因": "提取的根本原因",
    "纠正预防措施": ["措施1", "措施2", "措施3", "措施4"],
    "检测方法": "提取的检测方法",
    "溶液制备过程": "提取的溶液制备过程",
    "检测结果": "提取的检测结果",
    "初步原因分析": "提取的初步原因分析",
    "深入调查": "提取的深入调查内容",
    "对目前批次产品质量的影响": "提取的内容",
    "对相邻批次质量的影响": "提取的内容",
    "对其他产品质量或系统的影响": "提取的内容",
    "历史偏差": "回顾过去6个月内的重复性偏差情况",
    "其他影响": "其他方面影响",
    "调查结论": "提取的调查结论",
    "根本原因类别": ["选中的原因类别"],
    "CA_PA措施": [
        {{"措施": "措施1内容", "说明": "说明1"}},
        {{"措施": "措施2内容", "说明": "说明2"}},
        {{"措施": "措施3内容", "说明": "说明3"}},
        {{"措施": "措施4内容", "说明": "说明4"}}
    ]
}}
直接输出JSON：<|im_end|>
<|im_start|>assistant
"""

                from app.platform.ai.minimax_util import get_vision_util
                vision_util = get_vision_util()
                ai_result = await vision_util.recognize_text(prompt=prompt, max_tokens=8192)
                ai_data = _parse_ai_result(ai_result)
                logger.info(f"预览自动恢复成功，包含 {len(ai_data)} 个字段")

                # 更新数据库中的ai_result
                task.ai_result = ai_result
                task.update_time = datetime.now()
                await db.flush()

            except Exception as recovery_error:
                logger.error(f"自动恢复失败: {recovery_error}")
                # 如果自动恢复也失败，返回原始内容的HTML展示
                return ApiResponse(data={
                    'task_id': task_id,
                    'html_content': html_lib.escape(f'<div style="font-family:Times New Roman;font-size:12pt;padding:20px;">'
                        f'<p style="color:#666;">预览数据格式异常，已自动重新处理。</p>'
                        f'<pre style="white-space:pre-wrap;word-wrap:break-word;">{html_lib.escape(task.ai_result[:2000])}...</pre>'
                        f'</div>'),
                    'plain_content': f'<div><p>预览数据格式异常</p><pre>{task.ai_result[:2000]}...</pre></div>',
                })

        # 任务基本信息
        task_info = {
            'deviation_no': task.deviation_no,
            'creator': task.creator,
            'auditor': task.auditor,
            'report_date': task.report_date.isoformat() if task.report_date else '',
        }

        # 生成HTML预览
        html_content = _generate_preview_html(ai_data, task_info)

        # HTML转义用于JSON传输
        escaped_html = html_lib.escape(html_content)

        return ApiResponse(data={
            'task_id': task_id,
            'html_content': escaped_html,
            'plain_content': html_content,  # 未转义版本
        })

    except Exception as e:
        logger.error(f"生成预览失败: {e}")
        raise HTTPException(status_code=500, detail=f"生成预览失败: {str(e)}")


@router.post("/tasks/{task_id}/generate-standard", summary="生成标准Word文件")
async def generate_standard_docx(
    task_id: int,
    db: AsyncSession = Depends(get_db),
):
    """
    基于AI处理结果和模板文件生成标准格式的Word文档
    
    标准格式规范：
    - 纸张：A4纵向，上下边距1英寸，左2.29cm右1.91cm
    - 字体：Times New Roman + 宋体
    - 正文12pt，标题20pt
    - 行距：1.5倍行距
    - 对齐：标题居中，正文两端对齐
    """
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from app.platform.ai.minimax_util import get_vision_util
    import mammoth
    import json
    import copy

    # 获取任务
    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    if not task.ai_result:
        raise HTTPException(status_code=400, detail="请先进行AI处理")

    try:
        # 解析AI返回的JSON数据
        try:
            ai_data = _parse_ai_result(task.ai_result)
            logger.info(f"成功解析AI返回数据，包含 {len(ai_data)} 个字段")

        except (json.JSONDecodeError, ValueError) as e:
            # JSON解析失败，可能是HTML格式（用户编辑后保存的）
            # 自动重新读取原始文件并调用AI处理
            logger.warning(f"AI返回数据格式错误，自动重新处理: {e}")

            if not task.original_file_path:
                raise HTTPException(status_code=400, detail="原始文件不存在，无法重新处理")

            # 读取原始文件
            backend_dir = Path(__file__).resolve().parent.parent.parent.parent
            original_path = backend_dir / task.original_file_path.lstrip("/")

            if not original_path.exists():
                raise HTTPException(status_code=500, detail="原始文件不存在")

            with open(original_path, "rb") as f:
                text_result = mammoth.extract_raw_text(f)
                original_text = text_result.value

            # 获取SOP规则
            sop_result = await db.execute(select(SOPRule).where(SOPRule.status == 1))
            sop_rules = sop_result.scalars().all()
            sop_context = "\n\n".join([
                f"[{r.sop_code}] {r.sop_full_name}\n版本: {r.sop_version}\n业务标签: {r.business_tag or '无'}\n标准限值: {r.standard_limit or '无'}\n标准语句: {r.standard_sentence or '无'}\n"
                for r in sop_rules
            ])

            # 识别并加载SOP文件
            identified_sop_codes = _identify_sop_codes_in_text(original_text)
            sop_files_context = ""
            if identified_sop_codes:
                sop_files_context = "\n\n" + "="*60 + "\n相关 SOP 文件内容：\n" + "="*60 + "\n\n"
                for sop_code in identified_sop_codes:
                    sop_content = await _load_sop_file_content(sop_code, db, backend_dir)
                    if sop_content:
                        sop_files_context += f"【{sop_code}】\n{sop_content[:5000]}\n\n"

            # 调用AI处理
            prompt = f"""<|im_end|>
<|im_start|>user
你是一个专业的偏差报告处理助手。任务是：
1. 从用户上传的原始偏差报告中提取关键信息
2. 根据SOP规则库和相关SOP文件内容对内容进行标准化
3. 直接输出JSON格式的结构化数据

【关键要求】
- 不要输出任何思考过程、推理步骤或解释说明
- 不要使用<think>标签
- 不要使用省略号（...）或任何截断符号
- 所有文本内容必须完整输出，不要省略
- 直接输出JSON格式结果

SOP规则库：
{sop_context}
{sop_files_context}
标准报告结构模板：
{STANDARD_REPORT_TEMPLATE}

原始偏差报告内容：
{original_text}

请严格按照以下JSON格式输出，直接输出JSON：
{{
    "偏差描述": "提取的偏差描述内容",
    "根本原因": "提取的根本原因",
    "纠正预防措施": ["措施1", "措施2", "措施3", "措施4"],
    "检测方法": "提取的检测方法",
    "溶液制备过程": "提取的溶液制备过程",
    "检测结果": "提取的检测结果",
    "初步原因分析": "提取的初步原因分析",
    "深入调查": "提取的深入调查内容",
    "对目前批次产品质量的影响": "提取的内容",
    "对相邻批次质量的影响": "提取的内容",
    "对其他产品质量或系统的影响": "提取的内容",
    "历史偏差": "回顾过去6个月内的重复性偏差情况",
    "其他影响": "其他方面影响",
    "调查结论": "提取的调查结论",
    "根本原因类别": ["选中的原因类别"],
    "CA_PA措施": [
        {{"措施": "措施1内容", "说明": "说明1"}},
        {{"措施": "措施2内容", "说明": "说明2"}},
        {{"措施": "措施3内容", "说明": "说明3"}},
        {{"措施": "措施4内容", "说明": "说明4"}}
    ]
}}
直接输出JSON：<|im_end|>
<|im_start|>assistant
"""

            vision_util = get_vision_util()
            ai_result = await vision_util.recognize_text(prompt=prompt, max_tokens=8192)

            # 解析新的AI结果
            ai_data = _parse_ai_result(ai_result)
            logger.info(f"自动重新处理成功，包含 {len(ai_data)} 个字段")

            # 更新数据库中的ai_result
            task.ai_result = ai_result
            task.update_time = datetime.now()
            await db.flush()

        # 获取模板文件路径 - 从数据库获取启用的模板
        backend_dir = Path(__file__).resolve().parent.parent.parent.parent

        # 查询启用的模板
        template_result = await db.execute(
            select(ReportTemplate).where(ReportTemplate.is_active == 1)
        )
        active_template = template_result.scalar_one_or_none()

        if not active_template or not active_template.file_path:
            raise HTTPException(status_code=500, detail="未找到启用的报告模板，请在模板管理中启用模板")

        template_path = backend_dir / active_template.file_path.lstrip("/")

        if not template_path.exists():
            raise HTTPException(status_code=500, detail=f"模板文件不存在: {active_template.name}")

        # 打开模板文件（保留原始格式，包括字体、字号、边距等）
        doc = Document(str(template_path))

        # 【重要】不强制修改边距和字体设置，保留模板原始格式
        # 只进行占位符替换操作

        # 遍历文档中的所有段落，替换{{}}占位符
        for para in doc.paragraphs:
            text = para.text
            
            # 替换占位符
            if '{{偏差描述}}' in text or '{{偏差描述内容}}' in text:
                text = text.replace('{{偏差描述}}', ai_data.get('偏差描述', ''))
                text = text.replace('{{偏差描述内容}}', ai_data.get('偏差描述', ''))
            
            if '{{根本原因}}' in text or '{{根本原因内容}}' in text:
                text = text.replace('{{根本原因}}', ai_data.get('根本原因', ''))
                text = text.replace('{{根本原因内容}}', ai_data.get('根本原因', ''))
            
            if '{{措施1}}' in text:
                measures = ai_data.get('纠正预防措施', [])
                text = text.replace('{{措施1}}', measures[0] if len(measures) > 0 else '')
            
            if '{{措施2}}' in text:
                measures = ai_data.get('纠正预防措施', [])
                text = text.replace('{{措施2}}', measures[1] if len(measures) > 1 else '')
            
            if '{{措施3}}' in text:
                measures = ai_data.get('纠正预防措施', [])
                text = text.replace('{{措施3}}', measures[2] if len(measures) > 2 else '')
            
            if '{{措施4}}' in text:
                measures = ai_data.get('纠正预防措施', [])
                text = text.replace('{{措施4}}', measures[3] if len(measures) > 3 else '')
            
            if '{{检测方法内容}}' in text or '{{检测方法}}' in text:
                text = text.replace('{{检测方法内容}}', ai_data.get('检测方法', ''))
                text = text.replace('{{检测方法}}', ai_data.get('检测方法', ''))
            
            if '{{溶液制备过程}}' in text:
                text = text.replace('{{溶液制备过程}}', ai_data.get('溶液制备过程', ''))
            
            if '{{检测结果}}' in text:
                text = text.replace('{{检测结果}}', ai_data.get('检测结果', ''))
            
            if '{{初步原因分析}}' in text:
                text = text.replace('{{初步原因分析}}', ai_data.get('初步原因分析', ''))
            
            if '{{调查项}}' in text:
                text = text.replace('{{调查项}}', ai_data.get('深入调查', ''))
            
            if '{{调查内容}}' in text:
                text = text.replace('{{调查内容}}', ai_data.get('深入调查', ''))
            
            if '{{影响1}}' in text:
                text = text.replace('{{影响1}}', ai_data.get('对目前批次产品质量的影响', ''))
            
            if '{{影响2}}' in text:
                text = text.replace('{{影响2}}', ai_data.get('对相邻批次质量的影响', ''))
            
            if '{{影响3}}' in text:
                text = text.replace('{{影响3}}', ai_data.get('对其他产品质量或系统的影响', ''))
            
            if '{{历史偏差}}' in text:
                text = text.replace('{{历史偏差}}', ai_data.get('历史偏差', ''))
            
            if '{{其他影响}}' in text:
                text = text.replace('{{其他影响}}', ai_data.get('其他影响', ''))
            
            if '{{调查结论}}' in text:
                text = text.replace('{{调查结论}}', ai_data.get('调查结论', ''))
            
            if '{{其他原因}}' in text:
                root_causes = ai_data.get('根本原因类别', [])
                other_causes = [c for c in root_causes if '其他' in c]
                text = text.replace('{{其他原因}}', ','.join(other_causes) if other_causes else '')
            
            # 更新段落文本
            if text != para.text:
                # 保留原始格式，只更新文本
                for run in para.runs:
                    run.text = text
                    break  # 只更新第一个run

        # 处理表格中的占位符 - 智能替换每行
        ca_pa = ai_data.get('CA_PA措施', [])

        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        text = original_text

                        # 替换措施占位符
                        if '{{措施}}' in text:
                            # 获取当前行对应的措施
                            if '{{措施1}}' in text:
                                idx = 0
                            elif '{{措施2}}' in text:
                                idx = 1
                            elif '{{措施3}}' in text:
                                idx = 2
                            elif '{{措施4}}' in text:
                                idx = 3
                            else:
                                idx = 0
                            ca_item = ca_pa[idx] if idx < len(ca_pa) else {}
                            measure_text = ca_item.get('措施', '') if isinstance(ca_item, dict) else ''
                            text = text.replace('{{措施}}', measure_text)
                            text = text.replace('{{措施1}}', measure_text)
                            text = text.replace('{{措施2}}', '')
                            text = text.replace('{{措施3}}', '')
                            text = text.replace('{{措施4}}', '')

                        # 替换说明占位符
                        if '{{说明}}' in text:
                            if '{{措施1}}' in text:
                                idx = 0
                            elif '{{措施2}}' in text:
                                idx = 1
                            elif '{{措施3}}' in text:
                                idx = 2
                            elif '{{措施4}}' in text:
                                idx = 3
                            else:
                                idx = 0
                            ca_item = ca_pa[idx] if idx < len(ca_pa) else {}
                            desc_text = ca_item.get('说明', '') if isinstance(ca_item, dict) else ''
                            text = text.replace('{{说明}}', desc_text)

                        # 替换序号占位符
                        if '{{序号}}' in text:
                            text = text.replace('{{序号}}', str(row_idx + 1))

                        # 更新文本
                        if text != original_text:
                            for run in para.runs:
                                run.text = text
                                break

        # 保存文件
        output_dir = backend_dir / "uploads" / "deviation_automation"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_filename = f"standard_{task_id}_{uuid.uuid4().hex}.docx"
        output_path = output_dir / output_filename
        
        doc.save(str(output_path))
        
        # 更新任务记录
        relative_path = f"/uploads/deviation_automation/{output_filename}"
        task.standard_file_path = relative_path
        task.task_status = 4  # 已完成
        task.update_time = datetime.now()
        await db.flush()
        
        return ApiResponse(message="标准文件生成成功", data={
            "task_id": task_id,
            "standard_file_path": relative_path,
            "task_status": task.task_status,
        })
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"生成标准文件失败: {e}")
        raise HTTPException(status_code=500, detail=f"生成标准文件失败: {str(e)}")


# ============ 文件下载 ============

@router.get("/tasks/{task_id}/download/original", summary="下载原始文件")
async def download_original_file(
    task_id: int,
    db: AsyncSession = Depends(get_db),
):
    """下载原始Word文件"""
    from fastapi.responses import FileResponse

    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    if not task.original_file_path:
        raise HTTPException(status_code=404, detail="原始文件不存在")

    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    file_path = backend_dir / task.original_file_path.lstrip("/")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=str(file_path),
        filename=f"原始报告_{task.deviation_no}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.get("/tasks/{task_id}/download/standard", summary="下载标准报告文件")
async def download_standard_file(
    task_id: int,
    db: AsyncSession = Depends(get_db),
):
    """下载标准格式报告文件"""
    from fastapi.responses import FileResponse

    result = await db.execute(select(DevTask).where(DevTask.task_id == task_id))
    task = result.scalar_one_or_none()
    if not task:
        raise HTTPException(status_code=404, detail="任务不存在")

    if not task.standard_file_path:
        raise HTTPException(status_code=404, detail="标准报告文件不存在")

    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    file_path = backend_dir / task.standard_file_path.lstrip("/")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=str(file_path),
        filename=f"标准报告_{task.deviation_no}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.post("/sop-rules", summary="新增SOP规则")
async def create_sop_rule(
    rule: SOPRuleCreate,
    db: AsyncSession = Depends(get_db),
):
    """新增SOP规则"""
    # 检查编号唯一性
    result = await db.execute(
        select(SOPRule).where(SOPRule.sop_code == rule.sop_code)
    )
    existing = result.scalar_one_or_none()
    if existing:
        raise HTTPException(status_code=400, detail="该SOP编号已存在，请重新填写")

    new_rule = SOPRule(
        sop_code=rule.sop_code,
        sop_full_name=rule.sop_full_name,
        sop_version=rule.sop_version,
        business_tag=rule.business_tag,
        standard_limit=rule.standard_limit,
        standard_sentence=rule.standard_sentence,
        sop_file_path=rule.sop_file_path,
        status=1,
    )
    db.add(new_rule)
    await db.flush()
    await db.refresh(new_rule)

    return ApiResponse(message="新增成功", data={
        "id": new_rule.id,
        "sop_code": new_rule.sop_code,
    })


@router.get("/templates", summary="查询报告模板列表")
async def list_templates(
    is_active: Optional[int] = None,
    page: int = 1,
    page_size: int = 20,
    db: AsyncSession = Depends(get_db),
):
    """查询报告模板列表"""
    conditions = []
    if is_active is not None:
        conditions.append(ReportTemplate.is_active == is_active)

    query = select(ReportTemplate)
    if conditions:
        query = query.where(and_(*conditions))

    # 计数
    count_query = select(func.count()).select_from(query.subquery())
    count_result = await db.execute(count_query)
    total = count_result.scalar()

    # 分页
    query = query.order_by(ReportTemplate.create_time.desc())
    query = query.offset((page - 1) * page_size).limit(page_size)

    result = await db.execute(query)
    items = result.scalars().all()

    templates = []
    for t in items:
        templates.append({
            "id": t.id,
            "name": t.name,
            "description": t.description,
            "file_path": t.file_path,
            "is_active": t.is_active,
            "create_time": t.create_time.isoformat() if t.create_time else None,
        })

    return ApiResponse(data={"items": templates, "total": total})


class ReportTemplateCreate(BaseModel):
    name: str
    description: Optional[str] = None
    is_active: Optional[int] = 1


class ReportTemplateUpdate(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    file_path: Optional[str] = None
    is_active: Optional[int] = None


@router.post("/templates", summary="新增报告模板")
async def create_template(
    template: ReportTemplateCreate,
    db: AsyncSession = Depends(get_db),
):
    """新增报告模板"""
    new_template = ReportTemplate(
        name=template.name,
        description=template.description,
        is_active=template.is_active,
    )
    db.add(new_template)
    await db.flush()
    await db.refresh(new_template)

    return ApiResponse(message="新增成功", data={
        "id": new_template.id,
        "name": new_template.name,
    })


@router.put("/templates/{template_id}", summary="编辑报告模板")
async def update_template(
    template_id: int,
    template: ReportTemplateUpdate,
    db: AsyncSession = Depends(get_db),
):
    """编辑报告模板"""
    result = await db.execute(select(ReportTemplate).where(ReportTemplate.id == template_id))
    db_template = result.scalar_one_or_none()
    if not db_template:
        raise HTTPException(status_code=404, detail="模板不存在")

    update_data = template.model_dump(exclude_unset=True)
    for key, value in update_data.items():
        if value is not None:
            setattr(db_template, key, value)

    db_template.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="更新成功")


@router.post("/templates/upload", summary="上传模板文件")
async def upload_template_file(
    template_id: int = Form(...),
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
):
    """上传模板文件"""
    # 验证文件类型
    if not file.filename or not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="仅支持doc、docx格式文件")

    # 检查模板是否存在
    result = await db.execute(select(ReportTemplate).where(ReportTemplate.id == template_id))
    db_template = result.scalar_one_or_none()
    if not db_template:
        raise HTTPException(status_code=404, detail="模板不存在")

    # 保存文件
    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    uploads_dir = backend_dir / "uploads" / "templates"
    uploads_dir.mkdir(parents=True, exist_ok=True)

    file_ext = Path(file.filename).suffix.lower()
    saved_filename = f"template_{template_id}_{uuid.uuid4().hex}{file_ext}"
    file_path = uploads_dir / saved_filename

    content = await file.read()
    with open(file_path, "wb") as f:
        f.write(content)

    # 更新模板记录
    relative_path = f"/uploads/templates/{saved_filename}"
    db_template.file_path = relative_path
    db_template.update_time = datetime.now()
    await db.flush()

    return ApiResponse(message="上传成功", data={
        "template_id": template_id,
        "file_path": relative_path,
    })


@router.get("/templates/{template_id}/download", summary="下载模板文件")
async def download_template_file(
    template_id: int,
    db: AsyncSession = Depends(get_db),
):
    """下载模板文件"""
    from fastapi.responses import FileResponse

    result = await db.execute(select(ReportTemplate).where(ReportTemplate.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")

    if not template.file_path:
        raise HTTPException(status_code=404, detail="模板文件不存在")

    backend_dir = Path(__file__).resolve().parent.parent.parent.parent
    file_path = backend_dir / template.file_path.lstrip("/")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=str(file_path),
        filename=f"{template.name}.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@router.delete("/templates/{template_id}", summary="删除报告模板")
async def delete_template(
    template_id: int,
    db: AsyncSession = Depends(get_db),
):
    """删除报告模板"""
    result = await db.execute(select(ReportTemplate).where(ReportTemplate.id == template_id))
    template = result.scalar_one_or_none()
    if not template:
        raise HTTPException(status_code=404, detail="模板不存在")

    await db.delete(template)
    await db.flush()

    return ApiResponse(message="删除成功")


# ============ 数据库模型 ============

from sqlalchemy import Column, Integer, String, Text, SmallInteger, DateTime, Date
from sqlalchemy.orm import declarative_base

Base = declarative_base()


class SOPRule(Base):
    """SOP规则表"""
    __tablename__ = "sop_rule"
    __table_args__ = {"schema": "quality"}

    id = Column(Integer, primary_key=True, autoincrement=True)
    sop_code = Column(String(64), nullable=False, unique=True)
    sop_full_name = Column(String(256), nullable=False)
    sop_version = Column(String(32), nullable=False)
    business_tag = Column(String(256))
    standard_limit = Column(Text)
    standard_sentence = Column(Text)
    sop_file_path = Column(String(512))
    status = Column(SmallInteger, nullable=False, default=1)
    create_time = Column(DateTime, nullable=False, default=datetime.now)
    update_time = Column(DateTime, nullable=False, default=datetime.now, onupdate=datetime.now)


class DevTask(Base):
    """偏差任务表"""
    __tablename__ = "dev_task"
    __table_args__ = {"schema": "quality"}

    task_id = Column(Integer, primary_key=True, autoincrement=True)
    deviation_no = Column(String(64), nullable=False, unique=True)
    creator = Column(String(64), nullable=False)
    auditor = Column(String(64))
    report_date = Column(Date, nullable=False)
    original_file_path = Column(String(512))
    standard_file_path = Column(String(512))
    task_status = Column(SmallInteger, nullable=False, default=1)
    ai_result = Column(Text)
    create_time = Column(DateTime, nullable=False, default=datetime.now)
    update_time = Column(DateTime, nullable=False, default=datetime.now, onupdate=datetime.now)


class ReportTemplate(Base):
    """报告模板表"""
    __tablename__ = "report_template"
    __table_args__ = {"schema": "quality"}

    id = Column(Integer, primary_key=True, autoincrement=True)
    name = Column(String(128), nullable=False)
    description = Column(Text)
    file_path = Column(String(512))
    is_active = Column(SmallInteger, nullable=False, default=1)
    create_time = Column(DateTime, nullable=False, default=datetime.now)
    update_time = Column(DateTime, nullable=False, default=datetime.now, onupdate=datetime.now)
